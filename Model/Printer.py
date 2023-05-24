import cv2
import docx

import tkinter as tk
from tkinter import filedialog

from Model import Word
from Model import Paragraph


def write_document(paragraphs):
    document = docx.Document()

    for paragraph in paragraphs:
        write_paragraph(document, paragraph)

    save_document(document)


def write_paragraph(document, paragraph):
    p = document.add_paragraph()

    words = paragraph.get_words()

    style = document.styles['Normal']
    style.font.name = paragraph.get_font()

    for id, word in enumerate(words):
        if word.get_weight() == 'normal':
            p.add_run(word.get_text())
        elif word.get_weight() == 'bold':
            p.add_run(word.get_text()).bold = True
        elif word.get_weight() == 'italic':
            p.add_run(word.get_text()).italic = True
        elif word.get_weight() == 'underlined':
            p.add_run(word.get_text()).underlined = True
        else:
            print(f'there was an error writing the word {word.get_id()}:  {word.get_text()}')

        p.add_run(" ")


def save_dialog():
    path = filedialog.asksaveasfilename(defaultextension=".docx", initialfile="untitled.docx")

    if len(path) > 0:
        return path
    else:
        print("Save operation cancelled.")
        return None


def save_document(document):
    file_path = save_dialog()

    if file_path is not None:
        document.save(file_path)
